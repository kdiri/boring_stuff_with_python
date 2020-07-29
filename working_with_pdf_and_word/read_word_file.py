"""
.. module:: working_with_pdf_and_word.read_word_file
   :synopsis: Read word file.
"""
import docx
from loguru import logger

path = "pdf_word_files/demo.docx"


def get_text(file_name: str) -> str:
    doc = docx.Document(file_name)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return "\n".join(full_text)


def read_file():
    full_text = get_text(path)
    logger.success(full_text)
    return full_text


def create_doc():
    return docx.Document(path)


if __name__ == "__main__":
    read_file()
